import io
import streamlit as st
import os
import time
import tempfile
import unicodedata
from dotenv import load_dotenv
import fitz  # PyMuPDF
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
import chromadb
from groq import Groq
from fpdf import FPDF
from docx import Document as DocxDocument

load_dotenv()

# ─── Initialize Clients ───────────────────────────────────────────────────────

groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

@st.cache_resource
def load_embedding_model():
    return SentenceTransformer("all-MiniLM-L6-v2")

embedding_model = load_embedding_model()
chroma_client = chromadb.EphemeralClient()


# ─── Core Functions ───────────────────────────────────────────────────────────

def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text


def get_paper_title(pdf_path):
    doc = fitz.open(pdf_path)
    first_page = doc[0].get_text()
    return first_page[:500]


def chunk_text(text, chunk_size=500, chunk_overlap=50):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    return splitter.split_text(text)


def store_in_chromadb(chunks, collection_name):
    safe_name = collection_name.replace(" ", "_").replace("-", "_")[:50]
    try:
        chroma_client.delete_collection(safe_name)
    except Exception:
        pass
    collection = chroma_client.create_collection(
        name=safe_name,
        metadata={"hnsw:space": "cosine"}
    )
    embeddings = embedding_model.encode(chunks, show_progress_bar=False).tolist()
    collection.add(
        documents=chunks,
        embeddings=embeddings,
        ids=[f"chunk_{i}" for i in range(len(chunks))]
    )
    return collection


def retrieve_relevant_chunks(collection, query, n_results=5):
    query_embedding = embedding_model.encode([query]).tolist()
    results = collection.query(query_embeddings=query_embedding, n_results=n_results)
    return results["documents"][0]


def generate_answer(context, question, chat_history=None, temperature=0.3, answer_length=1024):
    system_prompt = f"""You are an expert research paper analyst. Use the following context from a research paper to answer the question.

Context:
{chr(10).join(context)}

Provide a detailed, accurate answer based only on the context provided. If the answer is not in the context, say "This information is not available in the provided paper."
"""
    messages = [{"role": "system", "content": system_prompt}]
    if chat_history:
        messages.extend(chat_history)
    else:
        messages.append({"role": "user", "content": question})

    response = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=messages,
        temperature=temperature,
        max_tokens=answer_length,
    )
    return response.choices[0].message.content


def generate_full_summary(collection, title_text, selected_sections, temperature, answer_length):
    all_questions = {
        "📄 Title & Authors": f"Based on this text from the first page, what is the title and who are the authors?\n\n{title_text}",
        "📌 One-Line Summary": "Summarize this research paper in one sentence.",
        "🎯 Problem Statement": "What problem does this paper try to solve?",
        "🔬 Methodology": "What methodology or approach did the authors use?",
        "📊 Key Results": "What are the key results and findings of this paper?",
        "⚠️ Limitations": "What are the limitations mentioned in this paper?",
        "🔮 Future Work": "What future work do the authors suggest?"
    }
    questions = {k: v for k, v in all_questions.items() if k in selected_sections}
    summary = {}
    for section, question in questions.items():
        chunks = retrieve_relevant_chunks(collection, question)
        answer = generate_answer(chunks, question, temperature=temperature, answer_length=answer_length)
        summary[section] = answer
    return summary


def clean_export_heading(section: str) -> str:
    if section and not section[0].isascii():
        return section.split(" ", 1)[-1]
    return section


def clean_pdf_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    return normalized.encode("latin-1", errors="replace").decode("latin-1")


# ─── Export Functions ─────────────────────────────────────────────────────────

def export_to_word(summary: dict, paper_name: str) -> bytes:
    doc = DocxDocument()
    doc.add_heading(f"Research Paper Summary: {paper_name}", 0)
    for section, content in summary.items():
        clean_section = clean_export_heading(section)
        doc.add_heading(clean_section, level=1)
        doc.add_paragraph(content)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_to_pdf(summary: dict, paper_name: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, clean_pdf_text(f"Research Paper Summary: {paper_name}"))
    pdf.ln(4)
    for section, content in summary.items():
        pdf.set_font("Helvetica", "B", 12)
        pdf.multi_cell(0, 8, clean_pdf_text(clean_export_heading(section)))
        pdf.set_font("Helvetica", size=10)
        safe_content = clean_pdf_text(content)
        pdf.multi_cell(0, 7, safe_content)
        pdf.ln(3)
    buffer = io.BytesIO()
    pdf.output(buffer)
    return buffer.getvalue()


# ─── Page Config ──────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Research Paper Summarizer",
    page_icon="📄",
    layout="wide"
)

# ─── Custom CSS ───────────────────────────────────────────────────────────────

st.markdown("""
<style>
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

section[data-testid="stSidebar"] {
    background-color: #0d0d1a;
    border-right: 1px solid #1e1e2e;
}
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3,
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] li,
section[data-testid="stSidebar"] label {
    color: #e2e2f0 !important;
}

div[data-testid="stButton"] button[kind="primary"] {
    background: linear-gradient(90deg, #6366f1, #8b5cf6);
    border: none;
    border-radius: 8px;
    font-weight: 600;
    transition: opacity 0.2s ease;
}
div[data-testid="stButton"] button[kind="primary"]:hover { opacity: 0.88; }
div[data-testid="stButton"] button[kind="secondary"] {
    border-radius: 8px;
    border: 1px solid #3d3d5c;
    font-weight: 500;
}

div[data-testid="stExpander"] {
    border: 1px solid #2a2a3e;
    border-radius: 12px;
    margin-bottom: 12px;
    background-color: #12121f;
}

div[data-testid="stMetric"] {
    background: #1a1a2e;
    border-radius: 12px;
    padding: 18px 20px;
    border: 1px solid #2d2d4e;
}
div[data-testid="stMetric"] label { color: #9090b0 !important; font-size: 13px !important; }
div[data-testid="stMetric"] div[data-testid="stMetricValue"] {
    font-size: 26px !important;
    font-weight: 700 !important;
    color: #a78bfa !important;
}

div[data-testid="stChatMessage"] {
    border-radius: 14px;
    padding: 6px 10px;
    margin-bottom: 10px;
    border: 1px solid #2a2a3e;
}

div[data-testid="stTabs"] button { font-weight: 600; font-size: 14px; }
div[data-testid="stDownloadButton"] button { border-radius: 8px; font-weight: 500; width: 100%; }
div[data-testid="stTextInput"] input {
    border-radius: 8px;
    border: 1px solid #3d3d5c;
    background-color: #12121f;
}

.section-card {
    background: #12121f;
    border: 1px solid #2a2a3e;
    border-radius: 12px;
    padding: 18px 20px;
    margin-bottom: 14px;
}
.section-card h4 { margin: 0 0 10px 0; color: #a78bfa; }
.section-card p  { margin: 0; color: #d0d0e8; line-height: 1.7; }
</style>
""", unsafe_allow_html=True)

# ─── Session State Init ───────────────────────────────────────────────────────

if "papers"         not in st.session_state: st.session_state["papers"]         = {}
if "chat_histories" not in st.session_state: st.session_state["chat_histories"] = {}

# ─── Sidebar ──────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown("## 📄 Research Summarizer")
    st.markdown("---")

    # ── Upload ──
    st.markdown("### 📂 Upload Papers")
    uploaded_files = st.file_uploader(
        "PDF files", type="pdf",
        accept_multiple_files=True,
        label_visibility="collapsed"
    )

    # ── Chunking Settings ──
    with st.expander("⚙️ Chunking Settings", expanded=False):
        chunk_size = st.slider("Chunk size (tokens)", 200, 1000, 500, step=50,
                               help="Larger = more context per chunk, fewer chunks total")
        chunk_overlap = st.slider("Chunk overlap", 0, 200, 50, step=10,
                                  help="Overlap between consecutive chunks")

    if uploaded_files:
        if st.button("🔍 Process All Papers", type="primary", use_container_width=True):
            progress = st.progress(0, text="Starting...")
            total = len(uploaded_files)
            for i, uploaded_file in enumerate(uploaded_files):
                paper_name = uploaded_file.name.replace(".pdf", "")
                if paper_name in st.session_state["papers"]:
                    continue
                progress.progress((i) / total, text=f"Processing: {paper_name}…")
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    tmp.write(uploaded_file.read())
                    tmp_path = tmp.name
                text       = extract_text_from_pdf(tmp_path)
                title_text = get_paper_title(tmp_path)
                chunks     = chunk_text(text, chunk_size, chunk_overlap)
                collection = store_in_chromadb(chunks, paper_name)
                os.unlink(tmp_path)
                st.session_state["papers"][paper_name] = {
                    "collection": collection,
                    "title_text": title_text,
                    "summary":    None,
                    "chunk_count": len(chunks)
                }
                st.session_state["chat_histories"][paper_name] = []
                progress.progress((i + 1) / total, text=f"Done: {paper_name}")
            progress.empty()
            st.success(f"✅ {len(st.session_state['papers'])} paper(s) ready!")

    # ── Loaded Papers ──
    if st.session_state["papers"]:
        st.markdown("---")
        st.markdown("### 📚 Loaded Papers")
        for name, data in st.session_state["papers"].items():
            badge = "✅" if data["summary"] else "⏳"
            st.markdown(f"{badge} **{name}**  \n`{data['chunk_count']} chunks`")

        st.markdown("")
        if st.button("🗑️ Clear All Papers", use_container_width=True):
            st.session_state["papers"]         = {}
            st.session_state["chat_histories"] = {}
            st.rerun()

    # ── AI Settings ──
    st.markdown("---")
    st.markdown("### 🤖 AI Settings")
    temperature = st.slider(
        "Creativity (temperature)", 0.0, 1.0, 0.3, step=0.05,
        help="Lower = more factual, higher = more creative"
    )
    answer_length = st.select_slider(
        "Answer length",
        options=[256, 512, 1024, 2048],
        value=1024,
        help="Max tokens in each answer"
    )
    n_chunks = st.slider(
        "Chunks to retrieve", 1, 10, 5,
        help="How many chunks are passed to the model as context"
    )

    st.markdown("---")
    st.markdown("### ⚙️ How it works")
    st.markdown("""
1. Upload one or more PDFs
2. Text extracted & chunked
3. Chunks embedded via **sentence-transformers**
4. Stored in **ChromaDB**
5. Semantic search retrieves top chunks
6. **LLaMA 3.3** generates answers
""")

# ─── Header ───────────────────────────────────────────────────────────────────

st.markdown("""
<div style='padding: 8px 0 20px 0'>
    <h1 style='margin-bottom: 4px'>📄 Research Paper Summarizer</h1>
    <p style='color: #9090b0; font-size: 16px; margin: 0'>
        Powered by RAG · ChromaDB · LLaMA 3.3 70B
    </p>
</div>
""", unsafe_allow_html=True)

# ─── Empty State ──────────────────────────────────────────────────────────────

if not st.session_state["papers"]:
    st.markdown("""
    <div style='text-align:center; padding: 80px 20px; color: #9090b0'>
        <div style='font-size: 56px; margin-bottom: 16px'>📄 → 🧠 → 📋</div>
        <h2 style='color: #c0c0d8; margin-bottom: 8px'>Welcome!</h2>
        <p style='font-size: 16px; max-width: 480px; margin: 0 auto'>
            Upload one or more research paper PDFs from the sidebar,
            then click <strong>Process All Papers</strong> to get started.
        </p>
        <br/>
        <p style='font-size: 13px; color: #555570'>
            Supports summarization · multi-turn chat · side-by-side comparison · PDF & Word export
        </p>
    </div>
    """, unsafe_allow_html=True)
    st.stop()

# ─── Stats Bar ────────────────────────────────────────────────────────────────

paper_names  = list(st.session_state["papers"].keys())
total_chunks = sum(p["chunk_count"] for p in st.session_state["papers"].values())
summarized   = sum(1 for p in st.session_state["papers"].values() if p["summary"])

c1, c2, c3, c4 = st.columns(4)
c1.metric("📄 Papers Loaded",  len(paper_names))
c2.metric("🧩 Total Chunks",   total_chunks)
c3.metric("✅ Summarized",     f"{summarized} / {len(paper_names)}")
c4.metric("🌡️ Temperature",    temperature)

st.markdown("<br/>", unsafe_allow_html=True)

# ─── Tabs ─────────────────────────────────────────────────────────────────────

tab_summary, tab_chat, tab_compare = st.tabs(
    ["📋 Summary & Export", "💬 Chat", "📊 Compare Papers"]
)


# ══════════════════════════════════════════════════════════════════════════════
# Tab 1 — Summary & Export
# ══════════════════════════════════════════════════════════════════════════════

with tab_summary:
    selected_paper = st.selectbox("Select a paper", paper_names, key="summary_select")
    paper_data     = st.session_state["papers"][selected_paper]

    # ── Section picker ──
    all_sections = [
        "📄 Title & Authors", "📌 One-Line Summary", "🎯 Problem Statement",
        "🔬 Methodology", "📊 Key Results", "⚠️ Limitations", "🔮 Future Work"
    ]
    selected_sections = st.multiselect(
        "Choose sections to include in summary",
        all_sections,
        default=all_sections,
        help="Deselect sections you don't need to speed things up"
    )

    col_gen, col_word, col_pdf = st.columns([2, 1, 1])

    with col_gen:
        generate_clicked = st.button(
            "🚀 Generate Summary", type="primary", use_container_width=True,
            disabled=not selected_sections
        )

    if generate_clicked and selected_sections:
        progress_bar = st.progress(0, text="Starting summary generation…")
        results      = {}
        all_q = {
            "📄 Title & Authors":  f"Based on this text from the first page, what is the title and who are the authors?\n\n{paper_data['title_text']}",
            "📌 One-Line Summary": "Summarize this research paper in one sentence.",
            "🎯 Problem Statement":"What problem does this paper try to solve?",
            "🔬 Methodology":      "What methodology or approach did the authors use?",
            "📊 Key Results":      "What are the key results and findings of this paper?",
            "⚠️ Limitations":      "What are the limitations mentioned in this paper?",
            "🔮 Future Work":      "What future work do the authors suggest?"
        }
        questions = {k: v for k, v in all_q.items() if k in selected_sections}
        total_q   = len(questions)
        start     = time.time()

        for idx, (section, question) in enumerate(questions.items()):
            progress_bar.progress(idx / total_q, text=f"Generating: {section}…")
            chunks          = retrieve_relevant_chunks(paper_data["collection"], question, n_chunks)
            results[section] = generate_answer(chunks, question, temperature=temperature, answer_length=answer_length)

        progress_bar.progress(1.0, text="Done!")
        time.sleep(0.4)
        progress_bar.empty()

        elapsed = round(time.time() - start, 1)
        st.session_state["papers"][selected_paper]["summary"] = results
        st.success(f"✅ Summary generated in {elapsed}s")
        st.rerun()

    summary = paper_data.get("summary")

    if summary:
        with col_word:
            st.download_button(
                "⬇️ Export Word",
                data=export_to_word(summary, selected_paper),
                file_name=f"{selected_paper}_summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        with col_pdf:
            st.download_button(
                "⬇️ Export PDF",
                data=export_to_pdf(summary, selected_paper),
                file_name=f"{selected_paper}_summary.pdf",
                mime="application/pdf",
                use_container_width=True
            )

        # ── Regenerate single section ──
        st.markdown("---")
        regen_section = st.selectbox(
            "🔄 Regenerate a single section",
            ["— select —"] + list(summary.keys()),
            key="regen_select"
        )
        if regen_section != "— select —":
            if st.button(f"🔄 Regenerate  {regen_section}"):
                all_q = {
                    "📄 Title & Authors":  f"Based on this text from the first page, what is the title and who are the authors?\n\n{paper_data['title_text']}",
                    "📌 One-Line Summary": "Summarize this research paper in one sentence.",
                    "🎯 Problem Statement":"What problem does this paper try to solve?",
                    "🔬 Methodology":      "What methodology or approach did the authors use?",
                    "📊 Key Results":      "What are the key results and findings of this paper?",
                    "⚠️ Limitations":      "What are the limitations mentioned in this paper?",
                    "🔮 Future Work":      "What future work do the authors suggest?"
                }
                with st.spinner(f"Regenerating {regen_section}…"):
                    chunks = retrieve_relevant_chunks(paper_data["collection"], all_q[regen_section], n_chunks)
                    new_ans = generate_answer(chunks, all_q[regen_section], temperature=temperature, answer_length=answer_length)
                st.session_state["papers"][selected_paper]["summary"][regen_section] = new_ans
                st.rerun()

        st.markdown("---")
        for section, content in summary.items():
            with st.expander(section, expanded=True):
                st.markdown(content)

    else:
        st.markdown("""
        <div style='text-align:center; padding: 40px; color: #9090b0'>
            <div style='font-size: 36px'>🧠</div>
            <p>Select sections above and click <strong>Generate Summary</strong>.</p>
        </div>
        """, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# Tab 2 — Chat
# ══════════════════════════════════════════════════════════════════════════════

with tab_chat:
    col_pick, col_info = st.columns([3, 1])
    with col_pick:
        selected_paper_chat = st.selectbox(
            "Select a paper to chat about", paper_names, key="chat_select"
        )
    with col_info:
        st.metric("💬 Messages", len(st.session_state["chat_histories"].get(selected_paper_chat, [])))

    collection   = st.session_state["papers"][selected_paper_chat]["collection"]
    chat_history = st.session_state["chat_histories"][selected_paper_chat]

    # ── Suggested Questions ──
    if not chat_history:
        st.markdown("#### 💡 Suggested questions — click to ask")
        suggestions = [
            "What is the main contribution of this paper?",
            "What dataset was used for evaluation?",
            "How does this compare to previous work?",
            "What are the key limitations?"
        ]
        s_cols = st.columns(2)
        for i, suggestion in enumerate(suggestions):
            with s_cols[i % 2]:
                if st.button(suggestion, key=f"suggestion_{i}", use_container_width=True):
                    chat_history.append({"role": "user", "content": suggestion})
                    chunks = retrieve_relevant_chunks(collection, suggestion, n_chunks)
                    answer = generate_answer(chunks, suggestion, chat_history=chat_history,
                                            temperature=temperature, answer_length=answer_length)
                    chat_history.append({"role": "assistant", "content": answer})
                    st.session_state["chat_histories"][selected_paper_chat] = chat_history
                    st.rerun()

        st.markdown("---")

    # ── Message history ──
    for msg in chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # ── Input ──
    if question := st.chat_input(f"Ask anything about '{selected_paper_chat}'…"):
        with st.chat_message("user"):
            st.markdown(question)
        chat_history.append({"role": "user", "content": question})

        with st.chat_message("assistant"):
            with st.spinner("Searching paper…"):
                chunks = retrieve_relevant_chunks(collection, question, n_chunks)
                answer = generate_answer(chunks, question, chat_history=chat_history,
                                         temperature=temperature, answer_length=answer_length)
            st.markdown(answer)

            with st.expander("📚 View retrieved chunks"):
                for i, chunk in enumerate(chunks):
                    st.caption(f"**Chunk {i+1}:** {chunk[:300]}…")

        chat_history.append({"role": "assistant", "content": answer})
        st.session_state["chat_histories"][selected_paper_chat] = chat_history

    # ── Clear ──
    if chat_history:
        if st.button("🗑️ Clear Chat History", key="clear_chat"):
            st.session_state["chat_histories"][selected_paper_chat] = []
            st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# Tab 3 — Compare Papers
# ══════════════════════════════════════════════════════════════════════════════

with tab_compare:
    if len(paper_names) < 2:
        st.markdown("""
        <div style='text-align:center; padding: 40px; color: #9090b0'>
            <div style='font-size: 36px'>📊</div>
            <p>Upload and process at least <strong>2 papers</strong> to use comparison.</p>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("### Compare papers side by side")

        selected_papers_to_compare = st.multiselect(
            "Select papers to compare",
            paper_names,
            default=paper_names[:2]
        )

        # ── Custom question compare ──
        compare_question = st.text_input(
            "Ask a question across all selected papers",
            placeholder="What dataset did the authors use?",
            key="compare_question"
        )

        # ── Quick compare buttons ──
        st.markdown("**Or pick a quick question:**")
        quick_cols = st.columns(4)
        quick_questions = [
            "What is the main contribution?",
            "What methodology was used?",
            "What are the key results?",
            "What are the limitations?"
        ]
        for i, q in enumerate(quick_questions):
            with quick_cols[i]:
                if st.button(q, key=f"quick_{i}", use_container_width=True):
                    st.session_state["quick_compare_q"] = q
                    st.rerun()

        if "quick_compare_q" in st.session_state and not compare_question:
            compare_question = st.session_state["quick_compare_q"]

        if st.button("🔍 Compare", type="primary") and compare_question and len(selected_papers_to_compare) >= 2:
            answers = {}
            with st.spinner("Searching all papers…"):
                for name in selected_papers_to_compare:
                    coll   = st.session_state["papers"][name]["collection"]
                    chunks = retrieve_relevant_chunks(coll, compare_question, n_chunks)
                    answer = generate_answer(chunks, compare_question,
                                             temperature=temperature, answer_length=answer_length)
                    answers[name] = {"answer": answer, "chunks": chunks}

            cols = st.columns(len(selected_papers_to_compare))
            for col, name in zip(cols, selected_papers_to_compare):
                with col:
                    st.markdown(f"**{name}**")
                    st.markdown(answers[name]["answer"])
                    with st.expander("📚 Source chunks"):
                        for i, chunk in enumerate(answers[name]["chunks"]):
                            st.caption(f"Chunk {i+1}: {chunk[:200]}…")

        st.markdown("---")
        st.markdown("### Section-by-section comparison")

        compare_sections = st.multiselect(
            "Choose sections to compare",
            ["📌 One-Line Summary", "🎯 Problem Statement", "🔬 Methodology",
             "📊 Key Results", "⚠️ Limitations", "🔮 Future Work"],
            default=["📌 One-Line Summary", "🎯 Problem Statement", "🔬 Methodology", "📊 Key Results"],
            key="compare_sections"
        )

        if st.button("📊 Generate Full Comparison", type="primary") and len(selected_papers_to_compare) >= 2:
            for section in compare_sections:
                st.markdown(f"#### {section}")
                section_answers = {}
                with st.spinner(f"Analyzing: {section}…"):
                    for name in selected_papers_to_compare:
                        coll   = st.session_state["papers"][name]["collection"]
                        chunks = retrieve_relevant_chunks(coll, section, n_chunks)
                        answer = generate_answer(chunks, section,
                                                 temperature=temperature, answer_length=answer_length)
                        section_answers[name] = answer

                cols = st.columns(len(selected_papers_to_compare))
                for col, name in zip(cols, selected_papers_to_compare):
                    with col:
                        st.markdown(f"**{name}**")
                        st.markdown(section_answers[name])
                st.markdown("---")
